from docx import Document
import fitz  # PyMuPDF
import requests
from bs4 import BeautifulSoup
import io

def read_docx(file):
    """Yüklenen .docx dosyasından metin okur"""
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def read_pdf(file):
    """Yüklenen .pdf dosyasından metin okur"""
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def read_source_pdf_from_url(pdf_url):
    """PDF URL'sinden sayfa sayfa metin okur"""
    response = requests.get(pdf_url)
    if response.status_code != 200:
        return []
    doc = fitz.open(stream=response.content, filetype="pdf")
    return [page.get_text() for page in doc]

def read_source_from_url(url):
    """
    PDF, DOCX veya HTML uzantılı URL'den metin okur.
    Sayfa/paragraf listesi döner.
    """
    if url.lower().endswith(".pdf"):
        return read_source_pdf_from_url(url)

    elif url.lower().endswith(".docx"):
        response = requests.get(url)
        doc = Document(io.BytesIO(response.content))
        return [para.text for para in doc.paragraphs if para.text.strip()]

    elif url.lower().endswith(".html") or "html" in url.lower():
        response = requests.get(url)
        soup = BeautifulSoup(response.text, "html.parser")
        text = soup.get_text()
        return [line.strip() for line in text.split("\n") if line.strip()]

    else:
        return []